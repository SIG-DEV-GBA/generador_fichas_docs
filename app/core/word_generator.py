"""
Generador de documentos Word desde fichas estructuradas.
"""

from pathlib import Path
from typing import Optional
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from loguru import logger

from app.models.ficha_schema import FichaData


class WordGenerator:
    """
    Genera documentos Word (.docx) desde FichaData estructurado.
    """

    def __init__(self, template_path: Optional[str | Path] = None):
        """
        Inicializa el generador de Word.

        Args:
            template_path: Ruta a plantilla .docx (opcional)
        """
        self.template_path = Path(template_path) if template_path else None

        if self.template_path and self.template_path.exists():
            logger.info(f"Usando plantilla: {self.template_path}")
        else:
            logger.info("Generando desde documento en blanco")

    def generate(
        self,
        ficha_data: FichaData,
        output_path: str | Path,
    ) -> Path:
        """
        Genera documento Word desde FichaData.

        Args:
            ficha_data: Datos de la ficha validados
            output_path: Ruta de salida del .docx

        Returns:
            Path al archivo generado
        """
        output_path = Path(output_path)
        logger.info(f"Generando Word: {output_path.name}")

        # Crear documento
        if self.template_path and self.template_path.exists():
            doc = Document(str(self.template_path))
        else:
            doc = Document()
            self._setup_styles(doc)

        # Añadir contenido
        self._add_header(doc, ficha_data)
        self._add_metadata(doc, ficha_data)
        self._add_main_content(doc, ficha_data)
        self._add_footer(doc, ficha_data)

        # Guardar
        doc.save(str(output_path))
        logger.info(f"✓ Documento generado: {output_path}")

        return output_path

    def _setup_styles(self, doc: Document) -> None:
        """Configura estilos del documento."""
        styles = doc.styles

        # Estilo para títulos de sección
        try:
            section_style = styles.add_style("SectionTitle", WD_STYLE_TYPE.PARAGRAPH)
            section_font = section_style.font
            section_font.name = "Arial"
            section_font.size = Pt(12)
            section_font.bold = True
            section_font.color.rgb = RGBColor(0, 0, 128)
        except:
            pass  # Estilo ya existe

    def _add_header(self, doc: Document, ficha: FichaData) -> None:
        """Añade encabezado del documento."""
        # Título principal
        title = doc.add_heading(ficha.nombre_ayuda, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # Espaciado

    def _add_metadata(self, doc: Document, ficha: FichaData) -> None:
        """Añade metadatos principales."""
        # Portales y categoría
        metadata_para = doc.add_paragraph()
        metadata_para.add_run("Portales: ").bold = True
        metadata_para.add_run(", ".join(ficha.portales))
        metadata_para.add_run(" | ")
        metadata_para.add_run("Categoría: ").bold = True
        metadata_para.add_run(", ".join(ficha.categoria))

        doc.add_paragraph()  # Espaciado

    def _add_main_content(self, doc: Document, ficha: FichaData) -> None:
        """Añade contenido principal de la ficha."""

        # Tipo de ayuda
        self._add_field(doc, "Tipo de ayuda", ficha.tipo_ayuda)

        # Fechas
        self._add_field(
            doc,
            "Fecha de inicio",
            ficha.fecha_inicio.strftime("%d/%m/%Y"),
        )
        self._add_field(
            doc,
            "Fecha de fin",
            ficha.fecha_fin.strftime("%d/%m/%Y"),
        )
        if ficha.fecha_publicacion:
            self._add_field(
                doc,
                "Fecha de publicación",
                ficha.fecha_publicacion.strftime("%d/%m/%Y"),
            )

        # Administración
        self._add_field(doc, "Ámbito territorial", ficha.ambito_territorial)
        self._add_field(doc, "Administración convocante", ficha.administracion)

        # Plazo
        self._add_field(doc, "Plazo de presentación", ficha.plazo_presentacion)

        # Beneficiarios y requisitos
        self._add_field(doc, "Beneficiarios/Destinatarios", ficha.beneficiarios, preserve_newlines=True)
        self._add_field(doc, "Requisitos de acceso", ficha.requisitos_acceso, preserve_newlines=True)

        # Descripción
        self._add_field(doc, "Descripción", ficha.descripcion, preserve_newlines=True)

        # Cuantía
        self._add_list_field(doc, "Cuantía", ficha.cuantia)
        self._add_field(doc, "Importe máximo", ficha.importe_maximo)

        # Resolución
        self._add_field(doc, "Resolución", ficha.resolucion, preserve_newlines=True)

        # Documentación
        self._add_list_field(doc, "Documentos a presentar", ficha.documentos_presentar)

        # Costes no subvencionables
        if ficha.costes_no_subvencionables:
            self._add_field(doc, "Costes no subvencionables", ficha.costes_no_subvencionables, preserve_newlines=True)

        # Criterios de concesión
        if ficha.criterios_concesion:
            self._add_field(doc, "Criterios de concesión", ficha.criterios_concesion, preserve_newlines=True)

        # Normativa
        self._add_list_field(doc, "Normativa reguladora", ficha.normativa_reguladora)
        if ficha.referencia_legislativa:
            self._add_list_field(doc, "Referencia legislativa", ficha.referencia_legislativa)

        # Lugar de presentación
        self._add_heading(doc, "Lugar y forma de presentación")
        if ficha.lugar_presentacion.presencial:
            p = doc.add_paragraph()
            p.add_run("Presencialmente en:\n").bold = True
            for lugar in ficha.lugar_presentacion.presencial:
                doc.add_paragraph(f"- {lugar}", style="List Bullet")

        if ficha.lugar_presentacion.electronica:
            p = doc.add_paragraph()
            p.add_run("Electrónicamente en:\n").bold = True
            for url in ficha.lugar_presentacion.electronica:
                doc.add_paragraph(f"- {url}", style="List Bullet")

    def _add_footer(self, doc: Document, ficha: FichaData) -> None:
        """Añade sección de Otros Datos."""
        doc.add_page_break()

        self._add_heading(doc, "Otros datos")

        self._add_field(doc, "USUARIO", ficha.otros_datos.USUARIO)
        self._add_field(doc, "FECHA", ficha.otros_datos.FECHA)

        if ficha.otros_datos.FRASE_PARA_PUBLICITAR:
            self._add_list_field(
                doc,
                "FRASE PARA PUBLICITAR",
                ficha.otros_datos.FRASE_PARA_PUBLICITAR,
            )

        if ficha.otros_datos.DOCUMENTOS_ADJUNTOS:
            self._add_list_field(
                doc,
                "DOCUMENTOS ADJUNTOS",
                ficha.otros_datos.DOCUMENTOS_ADJUNTOS,
            )

    def _add_heading(self, doc: Document, text: str) -> None:
        """Añade un encabezado de sección."""
        heading = doc.add_heading(text, level=2)
        heading.style = "Heading 2"

    def _add_field(
        self,
        doc: Document,
        label: str,
        value: str,
        preserve_newlines: bool = False,
    ) -> None:
        """
        Añade un campo con label en negrita.

        Args:
            doc: Documento
            label: Etiqueta del campo
            value: Valor del campo
            preserve_newlines: Si preservar saltos de línea
        """
        para = doc.add_paragraph()
        para.add_run(f"{label}: ").bold = True

        if preserve_newlines and "\n" in value:
            lines = value.split("\n")
            para.add_run(lines[0])
            for line in lines[1:]:
                para.add_run("\n" + line)
        else:
            para.add_run(value)

        para.space_after = Pt(6)

    def _add_list_field(self, doc: Document, label: str, items: list[str]) -> None:
        """
        Añade un campo con lista de items.

        Args:
            doc: Documento
            label: Etiqueta del campo
            items: Lista de elementos
        """
        para = doc.add_paragraph()
        para.add_run(f"{label}:").bold = True
        para.space_after = Pt(3)

        for item in items:
            doc.add_paragraph(item, style="List Bullet")

        doc.add_paragraph()  # Espaciado

    def generate_from_dict(
        self,
        ficha_dict: dict,
        output_path: str | Path,
    ) -> Path:
        """
        Genera Word desde diccionario (sin validar).

        Args:
            ficha_dict: Diccionario con datos de ficha
            output_path: Ruta de salida

        Returns:
            Path al archivo generado
        """
        ficha_data = FichaData(**ficha_dict)
        return self.generate(ficha_data, output_path)
